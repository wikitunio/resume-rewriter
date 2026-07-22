import streamlit as st
import pdfplumber
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from bs4 import BeautifulSoup
from groq import Groq
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from fpdf import FPDF
import io
import re

# --- Initialize Groq Client ---
try:
    api_key = st.secrets["GROQ_API_KEY"]
    client = Groq(api_key=api_key)
except KeyError:
    st.error("Missing GROQ_API_KEY. Please add it to your Streamlit Community Cloud Secrets.")
    st.stop()

# --- Streamlit Session State Initialization ---
if 'generation_complete' not in st.session_state:
    st.session_state.generation_complete = False
if 'final_resume' not in st.session_state:
    st.session_state.final_resume = ""
if 'final_cover_letter' not in st.session_state:
    st.session_state.final_cover_letter = ""
if 'review_feedback' not in st.session_state:
    st.session_state.review_feedback = ""
if 'new_score' not in st.session_state:
    st.session_state.new_score = 0.0

# --- Helper Functions ---
def sanitize_text(text):
    """Replaces Unicode characters (like non-breaking hyphens) with standard equivalents to prevent squished words."""
    replacements = {
        '•': '-', '–': '-', '—': '-', '‑': '-', '−': '-', # Catches all weird hyphens
        '‘': "'", '’': "'", '“': '"', '”': '"', '…': '...', '✓': 'v'
    }
    for search, replace in replacements.items():
        text = text.replace(search, replace)
    return text.encode('latin-1', 'ignore').decode('latin-1')

def extract_text_from_file(uploaded_file):
    text = ""
    file_extension = uploaded_file.name.split('.')[-1].lower()
    try:
        if file_extension == 'pdf':
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        elif file_extension == 'docx':
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif file_extension == 'txt':
            text = uploaded_file.read().decode('utf-8')
    except Exception as e:
        st.error(f"Error reading file: {e}")
    return text

def scrape_job_description(url):
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for element in soup(["script", "style", "header", "footer", "nav", "noscript", "svg"]):
            element.decompose()
        text = soup.get_text(separator="\n")
        lines = (line.strip() for line in text.splitlines())
        return "\n".join(chunk for chunk in lines if chunk)
    except Exception as e:
        st.error(f"Failed to extract job description from URL: {e}")
        return None

def calculate_ats_score(resume_text, job_description):
    if not resume_text.strip() or not job_description.strip():
        return 0.0
    vectorizer = TfidfVectorizer(stop_words='english')
    vectors = vectorizer.fit_transform([resume_text, job_description])
    similarity = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
    return round(similarity * 100, 2)

def optimize_resume(resume_text, job_description):
    prompt = f"""
    You are an expert ATS resume writer. Rewrite the provided resume to match the job description.
    Output ONLY the optimized resume in Markdown format.
    Job Description:\n{job_description}\n\nMy Resume:\n{resume_text}
    """
    completion = client.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
    )
    return completion.choices[0].message.content

def generate_cover_letter(resume_text, job_description):
    prompt = f"""
    Write a professional cover letter based on the applicant's resume and the job description.
    Output ONLY the cover letter text.
    Job Description:\n{job_description}\n\nMy Resume:\n{resume_text}
    """
    completion = client.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5,
    )
    return completion.choices[0].message.content

def finalize_documents(draft_resume, draft_cl, job_description):
    prompt = f"""
    You are an elite Executive Technical Recruiter. I am providing a Draft Resume and Draft Cover Letter. 

    Your task is to:
    1. Critique the drafts for impact.
    2. REWRITE BOTH documents to incorporate your critique, creating highly professional, natural-sounding final versions.
    
    CRITICAL DESIGNER INSTRUCTIONS FOR RESUME FORMAT:
    You MUST output the resume using this EXACT Markdown structure so my Python code can format it beautifully:
    
    # [Candidate Full Name]
    [City, Country] | [Phone] | [Email] | [LinkedIn]
    
    ## PROFESSIONAL SUMMARY
    [A strong 3-4 sentence paragraph]
    
    ## CORE COMPETENCIES
    - **[Skill Category]:** [Skill 1, Skill 2, Skill 3]
    
    ## PROFESSIONAL EXPERIENCE
    **[Job Title]** | [Company Name] | [Dates]
    - [Accomplishment 1]
    - [Accomplishment 2]

    Output your response exactly in this format using these strict delimiters (NO markdown code blocks):

    ===REVIEW===
    (Your concise explanation of the improvements made)
    ===FINAL_RESUME===
    (The finalized markdown resume)
    ===FINAL_COVER_LETTER===
    (The finalized cover letter)
    
    Job Description:\n{job_description}\n\nDraft Resume:\n{draft_resume}\n\nDraft Cover Letter:\n{draft_cl}
    """
    completion = client.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    
    response_text = completion.choices[0].message.content
    
    review_match = re.search(r'===REVIEW===(.*?)===FINAL_RESUME===', response_text, re.DOTALL)
    resume_match = re.search(r'===FINAL_RESUME===(.*?)===FINAL_COVER_LETTER===', response_text, re.DOTALL)
    cl_match = re.search(r'===FINAL_COVER_LETTER===(.*)', response_text, re.DOTALL)
    
    review = review_match.group(1).strip() if review_match else "Documents generated successfully."
    final_resume = resume_match.group(1).strip() if resume_match else draft_resume
    final_cl = cl_match.group(1).strip() if cl_match else draft_cl
    
    final_resume = final_resume.replace('```markdown', '').replace('```', '').strip()
    final_cl = final_cl.replace('```markdown', '').replace('```', '').strip()
    
    return review, final_resume, final_cl

# --- Designer-Grade PDF Generators ---
class DesignerPDF(FPDF):
    def chapter_title(self, title):
        self.set_font('helvetica', 'B', 12)
        self.set_text_color(33, 53, 71)
        self.cell(0, 6, title.upper(), 0, 1, 'L')
        self.line(self.get_x(), self.get_y(), 210 - self.get_x(), self.get_y())
        self.ln(3)
        self.set_text_color(0, 0, 0)

def create_pdf_from_markdown(md_text):
    pdf = DesignerPDF()
    pdf.add_page()
    pdf.set_margins(15, 15, 15)
    
    lines = md_text.split('\n')
    for i, line in enumerate(lines):
        line = sanitize_text(line.strip())
        if not line or re.match(r'^[-=*_]{3,}$', line):
            continue
            
        if line.startswith('# '):
            pdf.set_font("helvetica", "B", 16)
            pdf.cell(0, 8, line[2:].upper(), 0, 1, 'C')
        elif i > 0 and lines[i-1].strip().startswith('# '):
            pdf.set_font("helvetica", "", 10)
            pdf.cell(0, 5, line, 0, 1, 'C')
            pdf.ln(4)
        elif line.startswith('## '):
            pdf.ln(2)
            pdf.chapter_title(line[3:])
        elif line.startswith('- ') or line.startswith('* '):
            pdf.set_font("helvetica", "", 10)
            clean_line = line[2:].replace('**', '') 
            pdf.multi_cell(0, 5, "-  " + clean_line)
            pdf.ln(1)
        elif line.startswith('**') and ' | ' in line:
            pdf.set_font("helvetica", "B", 11)
            pdf.cell(0, 6, line.replace('**', ''), 0, 1, 'L')
        else:
            pdf.set_font("helvetica", "", 10)
            clean_line = line.replace('**', '')
            pdf.multi_cell(0, 5, clean_line)
            pdf.ln(1)
            
    return bytes(pdf.output())

def create_pdf_from_text(plain_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    pdf.set_font("helvetica", "", 11)
    
    for paragraph in plain_text.split('\n'):
        paragraph = sanitize_text(paragraph.strip())
        paragraph = paragraph.replace('**', '') # Strip bold asterisks for clean PDF output
        if paragraph and not re.match(r'^[-=*_]{3,}$', paragraph):
            pdf.multi_cell(0, 6, paragraph)
            pdf.ln(2)
            
    return bytes(pdf.output())

# --- Designer-Grade Word (.docx) Generators ---
def _add_formatted_runs(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)

def create_docx_from_markdown(md_text):
    doc = docx.Document()
    
    # Overwrite python-docx default author metadata
    doc.core_properties.author = "Waqar Ahmed Tunio"
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = md_text.split('\n')
    for i, line in enumerate(lines):
        line = line.strip()
        if not line or re.match(r'^[-=*_]{3,}$', line):
            continue
            
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].upper())
            run.font.size = Pt(16)
            run.bold = True
        elif i > 0 and lines[i-1].strip().startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(10)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].upper(), level=2)
            heading.runs[0].font.color.rgb = RGBColor(33, 53, 71)
            heading.runs[0].font.name = 'Calibri'
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, sanitize_text(line[2:]))
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, sanitize_text(line))
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_docx_from_text(plain_text):
    doc = docx.Document()
    
    # Overwrite python-docx default author metadata
    doc.core_properties.author = "Waqar Ahmed Tunio"
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    for paragraph in plain_text.split('\n'):
        paragraph = paragraph.strip()
        if paragraph and not re.match(r'^[-=*_]{3,}$', paragraph):
            p = doc.add_paragraph()
            _add_formatted_runs(p, sanitize_text(paragraph)) # Enables bolding in cover letter
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Streamlit Frontend ---
st.set_page_config(page_title="AI CV Optimizer", page_icon="📄", layout="wide")

st.title("CV Optimizer & Cover Letter Generator")
st.markdown("Powered by the ultra-fast **Groq API**")

option = st.radio("Choose Job Description Input Method:", ["Paste Job Description Text", "Paste Job Posting Link (URL)"])
job_desc = ""

if option == "Paste Job Description Text":
    job_desc = st.text_area("Paste Target Job Description:", height=150)
else:
    job_url = st.text_input("Paste Job Posting URL (e.g., https://...):")
    if job_url:
        with st.spinner("Scraping job description from URL..."):
            scraped_text = scrape_job_description(job_url)
            if scraped_text:
                job_desc = scraped_text
                st.success("Successfully retrieved job description from URL!")
                with st.expander("Preview Extracted Job Text"):
                    st.text(job_desc[:800] + "..." if len(job_desc) > 800 else job_desc)

uploaded_cv = st.file_uploader("Upload your current CV (.pdf, .docx, or .txt)", type=["pdf", "docx", "txt"])

if st.button("Optimize My CV & Generate Cover Letter", type="primary"):
    if not uploaded_cv or not job_desc:
        st.warning("Please upload a CV and provide a Job Description/URL first.")
    else:
        with st.spinner(f"Extracting text from {uploaded_cv.name}..."):
            raw_resume_text = extract_text_from_file(uploaded_cv)
            
        if not raw_resume_text.strip():
            st.error("Could not extract text from this file. It might be an empty file or a scanned image.")
        else:
            original_score = calculate_ats_score(raw_resume_text, job_desc)
            st.info(f"Baseline ATS Match Score: **{original_score}%**")
            
            try:
                with st.spinner("Agent 1: Drafting initial ATS resume..."):
                    draft_markdown = optimize_resume(raw_resume_text, job_desc)
                
                with st.spinner("Agent 2: Drafting human-sounding Cover Letter..."):
                    draft_cl = generate_cover_letter(raw_resume_text, job_desc)
                
                with st.spinner("Agent 3: Expert AI is rewriting documents for natural tone and ATS compliance..."):
                    review_feedback, final_resume, final_cover_letter = finalize_documents(draft_markdown, draft_cl, job_desc)
                
                st.session_state.final_resume = final_resume
                st.session_state.final_cover_letter = final_cover_letter
                st.session_state.review_feedback = review_feedback
                st.session_state.new_score = calculate_ats_score(final_resume, job_desc)
                st.session_state.original_score = original_score
                st.session_state.generation_complete = True
                
            except Exception as e:
                st.error(f"An error occurred during AI processing: {e}")

# --- Display Results from Session State ---
if st.session_state.generation_complete:
    st.success(f"Final ATS Match Score: **{st.session_state.new_score}%** (An improvement of {round(st.session_state.new_score - st.session_state.original_score, 2)}%)")
    
    st.markdown("---")
    
    tab1, tab2, tab3 = st.tabs(["📝 Final Tailored CV", "✉️ Final Cover Letter", "🔎 What The AI Improved"])
    
    with tab1:
        st.info("💡 **Tip:** Look for bracketed placeholders like `[Insert Number]` and manually replace them with real metrics before submitting!")
        st.markdown(st.session_state.final_resume)
        
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="⬇️ Download CV as PDF",
                data=create_pdf_from_markdown(st.session_state.final_resume),
                file_name="Optimized_ATS_Resume.pdf",
                mime="application/pdf"
            )
        with col2:
            st.download_button(
                label="⬇️ Download CV as Word (.docx)",
                data=create_docx_from_markdown(st.session_state.final_resume),
                file_name="Optimized_ATS_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
    with tab2:
        st.markdown(st.session_state.final_cover_letter)
        
        col3, col4 = st.columns(2)
        with col3:
            st.download_button(
                label="⬇️ Download Cover Letter as PDF",
                data=create_pdf_from_text(st.session_state.final_cover_letter),
                file_name="Cover_Letter.pdf",
                mime="application/pdf"
            )
        with col4:
            st.download_button(
                label="⬇️ Download Cover Letter as Word (.docx)",
                data=create_docx_from_text(st.session_state.final_cover_letter),
                file_name="Cover_Letter.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    with tab3:
        st.markdown("### The Expert Recruiter's Notes")
        st.markdown(st.session_state.review_feedback)
